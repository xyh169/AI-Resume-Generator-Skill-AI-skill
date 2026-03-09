import os
import sys
import subprocess
import time
from mcp.server.fastmcp import FastMCP

mcp = FastMCP(
    "Template to Markdown Converter",
    description="文档转换为 Markdown 的 AI 工具（带自动修复 & 重试 & 降级）",
)

MAX_RETRIES = 5
_log_buffer = []  # 收集日志供 MCP 返回


def _log(msg: str):
    _log_buffer.append(msg)


def _auto_install(package_name: str, pip_name: str | None = None) -> bool:
    pip_name = pip_name or package_name
    _log(f"[自动修复] 正在安装: {pip_name} ...")
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pip_name, "--quiet"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
        )
        _log(f"[自动修复] {pip_name} 安装成功。")
        return True
    except subprocess.CalledProcessError as e:
        _log(f"[自动修复] {pip_name} 安装失败: {e}")
        return False


def _ensure_pandoc() -> bool:
    try:
        import pypandoc
        pypandoc.get_pandoc_version()
        return True
    except Exception:
        _log("[自动修复] 正在自动下载 Pandoc ...")
        try:
            import pypandoc
            pypandoc.download_pandoc()
            _log("[自动修复] Pandoc 下载完成。")
            return True
        except Exception as e:
            _log(f"[自动修复] Pandoc 下载失败: {e}")
            return False


def _retry_with_fallbacks(file_path: str, strategies: list) -> str:
    """
    核心重试引擎。
    strategies: list of (策略名称, 转换函数, 修复函数)
    每个策略最多 MAX_RETRIES 次，失败切换下一策略。
    """
    all_errors = []

    for idx, (name, convert_fn, fix_fn) in enumerate(strategies):
        _log(f"[策略 {idx + 1}/{len(strategies)}] 尝试: {name}")

        for attempt in range(1, MAX_RETRIES + 1):
            try:
                result = convert_fn(file_path)
                _log(f"[成功] '{name}' 第 {attempt} 次尝试成功。")
                return result
            except ImportError as e:
                err_msg = f"  第 {attempt} 次: ImportError - {e}"
                _log(err_msg)
                all_errors.append(f"[{name}] {err_msg}")
                if fix_fn:
                    fix_fn()
                else:
                    break
            except Exception as e:
                err_msg = f"  第 {attempt} 次: {type(e).__name__} - {e}"
                _log(err_msg)
                all_errors.append(f"[{name}] {err_msg}")
                if attempt < MAX_RETRIES and fix_fn:
                    _log("  自动修复后重试 ...")
                    fix_fn()
                    time.sleep(0.3)
                elif attempt >= MAX_RETRIES:
                    _log(f"  策略 '{name}' 达到 {MAX_RETRIES} 次上限，切换。")
                    break

    error_report = "\n".join(all_errors)
    return (
        f"[错误] 所有转换策略均已失败 (文件: {file_path})。\n"
        f"建议: 请手动检查文件是否损坏，或手动安装相关依赖后重试。\n"
        f"--- 错误详情 ---\n{error_report}"
    )


# ============================================================
#  PDF 转换策略
# ============================================================
def _pdf_pymupdf4llm(fp):
    import pymupdf4llm
    return pymupdf4llm.to_markdown(fp)

def _pdf_pymupdf_raw(fp):
    import pymupdf
    doc = pymupdf.open(fp)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"## 第 {i+1} 页\n\n{text.strip()}")
    doc.close()
    if not pages:
        raise RuntimeError("PDF 无法提取到文本，可能是纯图片 PDF。")
    return "\n\n---\n\n".join(pages)

def _pdf_pdfplumber(fp):
    import pdfplumber
    pages = []
    with pdfplumber.open(fp) as pdf:
        for i, page in enumerate(pdf.pages):
            parts = []
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
            for table in page.extract_tables():
                if table:
                    header = table[0]
                    md = "| " + " | ".join(str(c or "") for c in header) + " |\n"
                    md += "| " + " | ".join("---" for _ in header) + " |\n"
                    for row in table[1:]:
                        md += "| " + " | ".join(str(c or "") for c in row) + " |\n"
                    parts.append(md)
            if parts:
                pages.append(f"## 第 {i+1} 页\n\n" + "\n\n".join(parts))
    if not pages:
        raise RuntimeError("pdfplumber 未提取到内容。")
    return "\n\n---\n\n".join(pages)

def _fix_pymupdf4llm(): _auto_install("pymupdf4llm")
def _fix_pymupdf(): _auto_install("pymupdf", "pymupdf")
def _fix_pdfplumber(): _auto_install("pdfplumber")


# ============================================================
#  DOCX 转换策略
# ============================================================
def _docx_python_docx(fp):
    import docx
    doc = docx.Document(fp)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            if para.style and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {para.text.strip()}")
            else:
                parts.append(para.text.strip())
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            md = "| " + " | ".join(rows[0]) + " |\n"
            md += "| " + " | ".join("---" for _ in rows[0]) + " |\n"
            for row in rows[1:]:
                md += "| " + " | ".join(row) + " |\n"
            parts.append(md)
    if not parts:
        raise RuntimeError("DOCX 文件未提取到内容。")
    return "\n\n".join(parts)

def _docx_pypandoc(fp):
    import pypandoc
    return pypandoc.convert_file(fp, 'md', format='docx')

def _fix_python_docx(): _auto_install("docx", "python-docx")
def _fix_pypandoc():
    _auto_install("pypandoc")
    _ensure_pandoc()


# ============================================================
#  TXT 转换策略
# ============================================================
def _txt_utf8(fp):
    with open(fp, 'r', encoding='utf-8') as f:
        return f.read()

def _txt_auto_detect(fp):
    import chardet
    with open(fp, 'rb') as f:
        raw = f.read()
    detected = chardet.detect(raw)
    enc = detected.get('encoding', 'utf-8')
    _log(f"  [自动检测] 编码: {enc} (置信度: {detected.get('confidence', 'N/A')})")
    return raw.decode(enc, errors='replace')

def _fix_chardet(): _auto_install("chardet")


# ============================================================
#  HTML / RTF / TEX 转换策略
# ============================================================
def _pandoc_convert(fmt):
    def _fn(fp):
        import pypandoc
        return pypandoc.convert_file(fp, 'md', format=fmt)
    return _fn

def _html_bs4_fallback(fp):
    from bs4 import BeautifulSoup
    with open(fp, 'r', encoding='utf-8', errors='replace') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    parts = []
    for tag in soup.find_all(['h1','h2','h3','h4','h5','h6','p','li']):
        if tag.name.startswith('h'):
            level = int(tag.name[1])
            parts.append(f"{'#' * level} {tag.get_text(strip=True)}")
        elif tag.name == 'li':
            parts.append(f"- {tag.get_text(strip=True)}")
        else:
            t = tag.get_text(strip=True)
            if t:
                parts.append(t)
    if not parts:
        parts.append(soup.get_text(separator='\n', strip=True))
    return "\n\n".join(parts)

def _fix_bs4(): _auto_install("bs4", "beautifulsoup4")


# ============================================================
#  MCP Tool 入口
# ============================================================

FORMAT_STRATEGIES = {
    '.pdf': [
        ("pymupdf4llm (高质量)", _pdf_pymupdf4llm, _fix_pymupdf4llm),
        ("PyMuPDF 原生提取", _pdf_pymupdf_raw, _fix_pymupdf),
        ("pdfplumber (含表格)", _pdf_pdfplumber, _fix_pdfplumber),
    ],
    '.docx': [
        ("python-docx (标题+表格)", _docx_python_docx, _fix_python_docx),
        ("pypandoc 降级", _docx_pypandoc, _fix_pypandoc),
    ],
    '.txt': [
        ("UTF-8 读取", _txt_utf8, None),
        ("chardet 自动编码检测", _txt_auto_detect, _fix_chardet),
    ],
    '.html': [
        ("pypandoc HTML→MD", _pandoc_convert('html'), _fix_pypandoc),
        ("BeautifulSoup 降级", _html_bs4_fallback, _fix_bs4),
    ],
    '.htm': [
        ("pypandoc HTML→MD", _pandoc_convert('html'), _fix_pypandoc),
        ("BeautifulSoup 降级", _html_bs4_fallback, _fix_bs4),
    ],
    '.rtf': [
        ("pypandoc RTF→MD", _pandoc_convert('rtf'), _fix_pypandoc),
    ],
    '.tex': [
        ("pypandoc LaTeX→MD", _pandoc_convert('latex'), _fix_pypandoc),
    ],
}


@mcp.tool()
def convert_file_to_md(file_path: str) -> str:
    """提取本地文档全部内容，将其格式化为 Markdown(md) 输出返回。
    支持 .pdf, .docx, .txt, .html, .rtf, .tex 等格式。
    具备自动安装依赖、错误重试（最多5次）、多策略降级转换能力。
    全部策略失败后会返回详细错误报告，建议手动处理。

    Args:
        file_path: 目标文档的绝对路径 (例如 /Users/.../resume.pdf)

    Returns:
        转换后的 Markdown 内容。如遇错误则返回包含修复过程的详细报告。
    """
    _log_buffer.clear()

    if not os.path.exists(file_path):
        return f"[错误] 文件不存在: {file_path}"

    ext = os.path.splitext(file_path)[1].lower()
    strategies = FORMAT_STRATEGIES.get(ext)

    if not strategies:
        supported = ', '.join(sorted(FORMAT_STRATEGIES.keys()))
        return f"[警告] 不支持的格式: {ext}。目前支持: {supported}"

    result = _retry_with_fallbacks(file_path, strategies)

    # 将日志附在结果之后（方便调试但不影响内容）
    if _log_buffer:
        log_text = "\n".join(_log_buffer)
        result = result + f"\n\n<!-- 转换日志\n{log_text}\n-->"

    return result


def main():
    mcp.run()


if __name__ == "__main__":
    main()
