import os
import sys
import subprocess
import argparse
import time
import traceback

# ============================================================
#  自动修复 & 重试核心机制
# ============================================================

MAX_RETRIES = 5  # 同一策略最大重试次数


def _auto_install(package_name, pip_name=None):
    """自动安装缺失的 Python 包，返回是否成功。"""
    pip_name = pip_name or package_name
    print(f"[自动修复] 正在安装缺失依赖: {pip_name} ...")
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pip_name, "--quiet"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
        )
        print(f"[自动修复] {pip_name} 安装成功。")
        return True
    except subprocess.CalledProcessError as e:
        print(f"[自动修复] {pip_name} 安装失败: {e}")
        return False


def _ensure_pandoc():
    """确保 pandoc 二进制可用，不可用则自动下载。"""
    try:
        import pypandoc
        pypandoc.get_pandoc_version()
        return True
    except Exception:
        print("[自动修复] 未检测到 Pandoc，正在自动下载 ...")
        try:
            import pypandoc
            pypandoc.download_pandoc()
            print("[自动修复] Pandoc 下载完成。")
            return True
        except Exception as e:
            print(f"[自动修复] Pandoc 下载失败: {e}")
            return False


def _retry_with_fallbacks(file_path, strategies):
    """
    核心重试引擎。
    strategies: list of (策略名称, 转换函数, 修复函数)
    每个策略最多尝试 MAX_RETRIES 次，失败则切换下一个策略。
    全部策略耗尽则返回错误报告。
    """
    all_errors = []

    for idx, (name, convert_fn, fix_fn) in enumerate(strategies):
        print(f"[策略 {idx + 1}/{len(strategies)}] 尝试: {name}")

        for attempt in range(1, MAX_RETRIES + 1):
            try:
                result = convert_fn(file_path)
                # 检查结果是否为错误字符串（兼容旧的返回方式）
                if isinstance(result, str) and result.startswith("[Error]"):
                    raise RuntimeError(result)
                print(f"[成功] 使用策略 '{name}' 在第 {attempt} 次尝试时成功。")
                return result
            except ImportError as e:
                err_msg = f"  第 {attempt} 次: ImportError - {e}"
                print(err_msg)
                all_errors.append(f"[{name}] {err_msg}")
                if fix_fn:
                    fix_fn()
                else:
                    break  # 没有修复函数，直接跳到下一策略
            except Exception as e:
                err_msg = f"  第 {attempt} 次: {type(e).__name__} - {e}"
                print(err_msg)
                all_errors.append(f"[{name}] {err_msg}")
                if attempt < MAX_RETRIES and fix_fn:
                    print(f"  正在尝试自动修复后重试 ...")
                    fix_fn()
                    time.sleep(0.5)
                elif attempt >= MAX_RETRIES:
                    print(f"  策略 '{name}' 已达 {MAX_RETRIES} 次上限，切换下一策略。")
                    break

    # 所有策略全部失败
    error_report = "\n".join(all_errors)
    return (
        f"[错误] 所有转换策略均已失败 (文件: {file_path})。\n"
        f"建议手动操作或检查文件是否损坏。\n"
        f"--- 错误详情 ---\n{error_report}"
    )


# ============================================================
#  各格式转换器 + 降级策略
# ============================================================

# ----- PDF -----
def _pdf_pymupdf4llm(file_path):
    import pymupdf4llm
    return pymupdf4llm.to_markdown(file_path)


def _pdf_pymupdf_raw(file_path):
    """降级策略：使用 PyMuPDF 原生提取纯文本"""
    import pymupdf  # pymupdf (fitz)
    doc = pymupdf.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"## 第 {i + 1} 页\n\n{text.strip()}")
    doc.close()
    if not pages:
        raise RuntimeError("PDF 无法提取到任何文本内容，可能是纯图片 PDF。")
    return "\n\n---\n\n".join(pages)


def _pdf_pdfplumber(file_path):
    """降级策略：使用 pdfplumber 提取文本和表格"""
    import pdfplumber
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            parts = []
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
            # 提取表格
            tables = page.extract_tables()
            for table in tables:
                if table:
                    # 转为 Markdown 表格
                    header = table[0]
                    md_table = "| " + " | ".join(str(c or "") for c in header) + " |\n"
                    md_table += "| " + " | ".join("---" for _ in header) + " |\n"
                    for row in table[1:]:
                        md_table += "| " + " | ".join(str(c or "") for c in row) + " |\n"
                    parts.append(md_table)
            if parts:
                pages.append(f"## 第 {i + 1} 页\n\n" + "\n\n".join(parts))
    if not pages:
        raise RuntimeError("pdfplumber 未能提取到内容。")
    return "\n\n---\n\n".join(pages)


def _fix_pdf_pymupdf4llm():
    _auto_install("pymupdf4llm")


def _fix_pdf_pymupdf():
    _auto_install("pymupdf", "pymupdf")


def _fix_pdf_pdfplumber():
    _auto_install("pdfplumber")


def convert_pdf(file_path):
    return _retry_with_fallbacks(file_path, [
        ("pymupdf4llm (高质量 Markdown)", _pdf_pymupdf4llm, _fix_pdf_pymupdf4llm),
        ("PyMuPDF 原生文本提取", _pdf_pymupdf_raw, _fix_pdf_pymupdf),
        ("pdfplumber (含表格提取)", _pdf_pdfplumber, _fix_pdf_pdfplumber),
    ])


# ----- DOCX -----
def _docx_python_docx(file_path):
    import docx
    doc = docx.Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            # 尝试保留标题层级
            if para.style and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {para.text.strip()}")
            else:
                parts.append(para.text.strip())
    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            md_table = "| " + " | ".join(rows[0]) + " |\n"
            md_table += "| " + " | ".join("---" for _ in rows[0]) + " |\n"
            for row in rows[1:]:
                md_table += "| " + " | ".join(row) + " |\n"
            parts.append(md_table)
    if not parts:
        raise RuntimeError("DOCX 文件未提取到任何内容。")
    return "\n\n".join(parts)


def _docx_pypandoc(file_path):
    """降级策略：通过 pypandoc 转换 DOCX"""
    import pypandoc
    return pypandoc.convert_file(file_path, 'md', format='docx')


def _fix_docx_python_docx():
    _auto_install("docx", "python-docx")


def _fix_docx_pypandoc():
    _auto_install("pypandoc")
    _ensure_pandoc()


def convert_docx(file_path):
    return _retry_with_fallbacks(file_path, [
        ("python-docx (含标题+表格)", _docx_python_docx, _fix_docx_python_docx),
        ("pypandoc 降级转换", _docx_pypandoc, _fix_docx_pypandoc),
    ])


# ----- TXT -----
def _txt_utf8(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def _txt_auto_detect(file_path):
    """降级：自动检测编码"""
    import chardet
    with open(file_path, 'rb') as f:
        raw = f.read()
    detected = chardet.detect(raw)
    encoding = detected.get('encoding', 'utf-8')
    print(f"  [自动检测] 编码: {encoding} (置信度: {detected.get('confidence', 'N/A')})")
    return raw.decode(encoding, errors='replace')


def _fix_txt_chardet():
    _auto_install("chardet")


def convert_txt(file_path):
    return _retry_with_fallbacks(file_path, [
        ("UTF-8 直接读取", _txt_utf8, None),
        ("chardet 自动检测编码", _txt_auto_detect, _fix_txt_chardet),
    ])


# ----- HTML / RTF / TEX (pypandoc 系列) -----
def _pandoc_convert(format_type):
    def _fn(file_path):
        import pypandoc
        return pypandoc.convert_file(file_path, 'md', format=format_type)
    return _fn


def _html_bs4_fallback(file_path):
    """降级：使用 BeautifulSoup 做 HTML 到文本的粗提取"""
    from bs4 import BeautifulSoup
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    # 提取标题
    parts = []
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
        if tag.name.startswith('h'):
            level = int(tag.name[1])
            parts.append(f"{'#' * level} {tag.get_text(strip=True)}")
        elif tag.name == 'li':
            parts.append(f"- {tag.get_text(strip=True)}")
        else:
            text = tag.get_text(strip=True)
            if text:
                parts.append(text)
    if not parts:
        # 最后兜底：全文本
        parts.append(soup.get_text(separator='\n', strip=True))
    return "\n\n".join(parts)


def _fix_pypandoc():
    _auto_install("pypandoc")
    _ensure_pandoc()


def _fix_bs4():
    _auto_install("bs4", "beautifulsoup4")


def convert_html(file_path):
    return _retry_with_fallbacks(file_path, [
        ("pypandoc HTML→MD", _pandoc_convert('html'), _fix_pypandoc),
        ("BeautifulSoup 降级提取", _html_bs4_fallback, _fix_bs4),
    ])


def convert_rtf(file_path):
    return _retry_with_fallbacks(file_path, [
        ("pypandoc RTF→MD", _pandoc_convert('rtf'), _fix_pypandoc),
    ])


def convert_tex(file_path):
    return _retry_with_fallbacks(file_path, [
        ("pypandoc LaTeX→MD", _pandoc_convert('latex'), _fix_pypandoc),
    ])


# ============================================================
#  主调度
# ============================================================

FORMAT_MAP = {
    '.pdf': convert_pdf,
    '.docx': convert_docx,
    '.txt': convert_txt,
    '.html': convert_html,
    '.htm': convert_html,
    '.rtf': convert_rtf,
    '.tex': convert_tex,
}


def process_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    converter = FORMAT_MAP.get(ext)
    if converter:
        return converter(file_path)
    else:
        return f"[警告] 不支持的格式: {ext}，目前支持: {', '.join(FORMAT_MAP.keys())}"


def main():
    parser = argparse.ArgumentParser(
        description="转换多种格式模板文档至 Markdown 格式（带自动修复 & 重试 & 降级）。"
    )
    parser.add_argument("-i", "--input", required=True, help="输入文件或文件夹路径")
    parser.add_argument("-o", "--output", required=True, help="Markdown 输出文件夹")
    args = parser.parse_args()

    input_path = args.input
    output_dir = args.output

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    files_to_process = []
    if os.path.isfile(input_path):
        files_to_process.append(input_path)
    elif os.path.isdir(input_path):
        for filename in sorted(os.listdir(input_path)):
            fp = os.path.join(input_path, filename)
            if os.path.isfile(fp):
                files_to_process.append(fp)
    else:
        print(f"[错误] 路径不存在: {input_path}")
        sys.exit(1)

    total = len(files_to_process)
    success = 0
    failed = []

    for i, fp in enumerate(files_to_process):
        filename = os.path.basename(fp)
        print(f"\n{'='*60}")
        print(f"[{i+1}/{total}] 解析中: {filename}")
        print(f"{'='*60}")

        md_content = process_file(fp)

        base_name = os.path.splitext(filename)[0]
        output_path = os.path.join(output_dir, f"{base_name}.md")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

        if md_content.startswith("[错误]"):
            failed.append(filename)
            print(f"[失败] {filename} -> {output_path} (含错误信息)")
        else:
            success += 1
            print(f"[完成] {filename} -> {output_path}")

    # 汇总报告
    print(f"\n{'='*60}")
    print(f"转换完成: 成功 {success}/{total}")
    if failed:
        print(f"失败文件: {', '.join(failed)}")
        print("建议: 请检查以上文件是否损坏，或手动安装对应依赖后重试。")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
